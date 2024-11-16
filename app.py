import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os

# Ruta de la imagen de la firma (reemplaza con la ruta correcta en tu sistema)
firma_imagen_ruta = 'firma.jpg'

# Verificar si la imagen de firma existe
if not os.path.exists(firma_imagen_ruta):
    st.error("La imagen de la firma no se encuentra en el directorio.")
else:
    # Título de la página
    st.title("Generar Documento Word con Firma")

    # Descripción
    st.write("Ingresa tu nombre y haz clic en el botón para generar el documento Word con tu firma.")

    # Campo de texto para ingresar el nombre
    nombre = st.text_input("Nombre:", placeholder="Ingrese su nombre")

    # Botón para generar y descargar el Word
    if st.button("Generar y Descargar Word"):
        if not nombre.strip():
            st.error("Por favor, ingresa un nombre.")
        else:
            # Crear un documento Word
            doc = Document()

            # Agregar texto al documento
            doc.add_paragraph(f"Este documento es del aprendiz: {nombre}.")

            # Insertar la imagen de la firma
            try:
                doc.add_paragraph("Firma del instructor:")
                doc.add_picture(firma_imagen_ruta, width=Inches(2))  # Ajusta el tamaño de la imagen si es necesario
            except Exception as e:
                st.error(f"Error al insertar la firma: {str(e)}")

            # Agregar línea para la firma del instructor
            doc.add_paragraph("Firma del instructor: _______________")

            # Guardar el archivo Word en un buffer en memoria
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Proporcionar el archivo para descarga
            st.download_button(
                label="Descargar Documento Word",
                data=buffer,
                file_name="firma_documento.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Función para verificar competencias (simulación)
    def verificar_competencias(ficha_id):
        # Simulación de respuesta del servidor
        return ficha_id in links and ficha_id == "2876293"  # Por ejemplo, solo "2876293" está completo

    # Diccionario de IDs de ficha con sus URLs correspondientes
    links = {
        "2876293": "https://1drv.ms/x/c/7f0aa27d6eb5cc8f/EVm8_aNEU-1HklqlpaEjqEEBCdR2uXaVwdCzcZhgbPV9Pg?e=DDU7ZC",
        "2368753": "https://1drv.ms/x/c/7f0aa27d6eb5cc8f/Ee2vuJfHdZJLvRTDo-l0BhcBOnzGvKszPeCI1OTo_ta4wQ?e=6lOpEM",
        # Agrega más IDs de ficha y enlaces según sea necesario
    }

    # Descripción de la página para redirigir a Excel
    st.write("Ingrese su número de ficha para redirigir al archivo Excel asociado:")

    # Campo de entrada para el número de ficha
    ficha_id = st.text_input("Número de ficha:", placeholder="Ingrese el número de ficha")

    # Botón para redirigir al enlace del Excel
    if st.button("Redirigir a su Excel de Ficha"):
        if ficha_id in links:
            st.success(f"Redirigiendo a Excel de la ficha {ficha_id}...")
            st.markdown(f"[Abrir Excel de Ficha {ficha_id}]({links[ficha_id]})", unsafe_allow_html=True)
        else:
            st.error("Ficha no encontrada. Por favor, verifica el número de ficha.")

    # Botón para verificar competencias
    if st.button("Verificar Competencias"):
        if ficha_id:
            completo = verificar_competencias(ficha_id)
            if completo:
                st.success("Todas las competencias están completas. Documento Word generado.")
            else:
                st.warning("Competencias incompletas, no se generará el documento Word.")
        else:
            st.error("Por favor, ingrese un número de ficha para verificar competencias.")

    # Pie de página
    st.write("¿Necesitas ayuda? [Contacta con soporte](#)")
